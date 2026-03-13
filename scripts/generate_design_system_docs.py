from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


OUT_DIR = Path("output/doc")
PDF_PATH = OUT_DIR / "H2_Design_System.pdf"
DOCX_PATH = OUT_DIR / "H2_Design_System.docx"

BG = "#F6F7F1"
TEXT = "#111111"
MUTED = "#555555"
SOFT = "#888888"
LABEL = "#AAAAAA"
OLIVE = "#ADB35B"
CARD = "#E7E3D3"

COLOR_TOKENS = [
    ("Background", BG, "Primary page canvas"),
    ("Text", TEXT, "Main headings and body anchors"),
    ("Muted", MUTED, "Secondary copy"),
    ("Tertiary", SOFT, "Section metadata and subdued UI"),
    ("Label", LABEL, "Overlines and utility labels"),
    ("Accent Olive", OLIVE, "Progress and highlight accent"),
    ("Card", CARD, "Project card surface"),
]

TYPE_ROLES = [
    ("Section Label", "12px, uppercase, wide tracking, medium weight"),
    ("Display Heading", "40px to 112px, bold, tight tracking"),
    ("Narrative Serif", "48px editorial statement, relaxed line height"),
    ("Body Copy", "14px to 20px, neutral rhythm, muted tone"),
    ("Utility Text", "10px to 12px, uppercase, system-like spacing"),
]

COMPONENT_PATTERNS = [
    ("Hero", "Full-viewport, dark cinematic stage with image treatment and a single centered statement."),
    ("Narrative", "Sticky editorial reading panel with serif-led messaging and an outlined CTA."),
    ("Projects", "Radial gallery of dark image cards, high contrast type, minimal metadata chips."),
    ("Timeline", "Horizontal pinned scroll with olive progress line and oversized ghost typography."),
    ("Contact", "Oversized type-led section with minimal chrome and direct contact actions."),
    ("Footer", "Thin top rule, minimal metadata, decorative cat element aligned to the divider."),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_page_background(section, fill: str) -> None:
    sect_pr = section._sectPr
    pg_borders = sect_pr.first_child_found_in("w:pgBorders")
    if pg_borders is not None:
        sect_pr.remove(pg_borders)
    background = OxmlElement("w:background")
    background.set(qn("w:color"), fill.replace("#", ""))
    sect_pr.append(background)


def build_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Avenir Next"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string(TEXT.replace("#", ""))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("H2 Personal Website\nDesign System")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(TEXT.replace("#", ""))

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("A concise visual and interaction system derived from the production site.")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(SOFT.replace("#", ""))

    document.add_paragraph("")

    intro = document.add_paragraph()
    intro.add_run("Direction: ").bold = True
    intro.add_run("editorial, minimal, monochrome-first, soft off-white canvas, sharp black typography, sparse olive accent.")

    heading = document.add_paragraph()
    run = heading.add_run("Color System")
    run.bold = True
    run.font.size = Pt(16)

    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Token", "Hex", "Usage"]
    for idx, label in enumerate(headers):
      cell = table.rows[0].cells[idx]
      cell.text = label
      set_cell_shading(cell, CARD)
    for name, hex_value, usage in COLOR_TOKENS:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = hex_value
        row[2].text = usage
        set_cell_shading(row[1], hex_value)

    heading = document.add_paragraph()
    run = heading.add_run("Typography")
    run.bold = True
    run.font.size = Pt(16)

    for role, spec in TYPE_ROLES:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(f"{role}: ").bold = True
        p.add_run(spec)

    heading = document.add_paragraph()
    run = heading.add_run("Component Patterns")
    run.bold = True
    run.font.size = Pt(16)

    for name, description in COMPONENT_PATTERNS:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(f"{name}: ").bold = True
        p.add_run(description)

    heading = document.add_paragraph()
    run = heading.add_run("Implementation Notes")
    run.bold = True
    run.font.size = Pt(16)

    notes = [
        "Primary background is #F6F7F1 across light sections and footer.",
        "Black text and thin black borders establish structure before color is introduced.",
        "Olive (#ADB35B) should remain sparse and signal motion, status, or emphasis.",
        "Motion should feel cinematic and deliberate, not playful or noisy.",
        "Section labels should remain small, uppercase, and low-contrast.",
    ]
    for note in notes:
        document.add_paragraph(note, style="List Bullet")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(DOCX_PATH)


def build_pdf() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title="H2 Personal Website Design System",
    )

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="TitleLarge",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=28,
            textColor=colors.HexColor(TEXT),
            alignment=TA_CENTER,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Subtitle",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=11,
            leading=15,
            textColor=colors.HexColor(SOFT),
            alignment=TA_CENTER,
            spaceAfter=24,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Section",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=18,
            textColor=colors.HexColor(TEXT),
            spaceBefore=12,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodySmall",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=15,
            textColor=colors.HexColor(MUTED),
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BulletSmall",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=15,
            textColor=colors.HexColor(MUTED),
            leftIndent=12,
            bulletIndent=0,
            spaceAfter=4,
        )
    )

    story = [
        Paragraph("H2 Personal Website Design System", styles["TitleLarge"]),
        Paragraph(
            "A production-grounded visual system covering color, typography, layout rhythm, interaction patterns, and component rules.",
            styles["Subtitle"],
        ),
        Paragraph(
            "Direction: editorial, minimal, monochrome-first, soft off-white canvas, sharp black typography, muted gray support text, and a sparse olive accent.",
            styles["BodySmall"],
        ),
        Spacer(1, 8),
        Paragraph("Color System", styles["Section"]),
    ]

    color_rows = [["Token", "Hex", "Usage"]] + [list(row) for row in COLOR_TOKENS]
    color_table = Table(color_rows, colWidths=[110, 80, 280])
    color_style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(CARD)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(TEXT)),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("LEADING", (0, 0), (-1, -1), 13),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D7D9D1")),
        ("BACKGROUND", (1, 1), (1, 1), colors.HexColor(BG)),
        ("BACKGROUND", (1, 2), (1, 2), colors.HexColor(TEXT)),
        ("TEXTCOLOR", (1, 2), (1, 2), colors.white),
        ("BACKGROUND", (1, 3), (1, 3), colors.HexColor(MUTED)),
        ("TEXTCOLOR", (1, 3), (1, 3), colors.white),
        ("BACKGROUND", (1, 4), (1, 4), colors.HexColor(SOFT)),
        ("TEXTCOLOR", (1, 4), (1, 4), colors.white),
        ("BACKGROUND", (1, 5), (1, 5), colors.HexColor(LABEL)),
        ("BACKGROUND", (1, 6), (1, 6), colors.HexColor(OLIVE)),
        ("BACKGROUND", (1, 7), (1, 7), colors.HexColor(CARD)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor(BG), colors.white]),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]
    color_table.setStyle(TableStyle(color_style))
    story.extend([color_table, Spacer(1, 10), Paragraph("Typography", styles["Section"])])

    for role, spec in TYPE_ROLES:
        story.append(Paragraph(f"<b>{role}</b>: {spec}", styles["BulletSmall"], bulletText="-"))

    story.extend([Spacer(1, 8), Paragraph("Component Patterns", styles["Section"])])
    for name, description in COMPONENT_PATTERNS:
        story.append(Paragraph(f"<b>{name}</b>: {description}", styles["BulletSmall"], bulletText="-"))

    story.extend([Spacer(1, 8), Paragraph("Implementation Notes", styles["Section"])])
    for note in [
        "Light sections use a shared off-white canvas rather than pure white.",
        "Black structure, thin borders, and generous spacing create the core visual tone.",
        "Olive should remain sparse and purposeful.",
        "Motion should feel cinematic and paced, not playful.",
        "Most labels are uppercase with wide tracking and low contrast.",
    ]:
        story.append(Paragraph(note, styles["BulletSmall"], bulletText="-"))

    def draw_bg(canvas, _doc):
        canvas.saveState()
        canvas.setFillColor(colors.HexColor(BG))
        canvas.rect(0, 0, A4[0], A4[1], fill=1, stroke=0)
        canvas.restoreState()

    doc.build(story, onFirstPage=draw_bg, onLaterPages=draw_bg)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(PDF_PATH.resolve())
    print(DOCX_PATH.resolve())
